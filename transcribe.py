import os
import subprocess
import argparse
import re
import json
import threading
import queue
import signal
import sys
import nltk
import torch
from datetime import datetime
from docx import Document
from openpyxl import load_workbook
from pathlib import Path
from faster_whisper import WhisperModel, BatchedInferencePipeline

nltk.download("punkt", quiet=True)
nltk.download("punkt_tab", quiet=True)
from nltk.tokenize import sent_tokenize


# -- Graceful shutdown --------------------------------------------------------
# A single global Event. Any part of the code can check _stop.is_set() to
# know that the user pressed CTRL+C and we should stop after the current task.

_stop = threading.Event()

def _handle_sigint(sig, frame):
    if not _stop.is_set():
        print("\n\n  [!] CTRL+C detected -- finishing current transcription then stopping.")
        print("      Press CTRL+C again to force-quit immediately.\n")
        _stop.set()
    else:
        print("\n  [!] Force-quitting.")
        sys.exit(1)

signal.signal(signal.SIGINT, _handle_sigint)


# -- Helpers ------------------------------------------------------------------

def format_timestamp(seconds: float) -> str:
    seconds = int(seconds)
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    return f"{h:02}:{m:02}:{s:02}"

def format_duration(seconds: float) -> str:
    seconds = int(seconds)
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    parts = []
    if h:
        parts.append(f"{h}h")
    if m:
        parts.append(f"{m}m")
    parts.append(f"{s}s")
    return " ".join(parts)

def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "", name).strip()

def get_video_metadata(url: str) -> dict | None:
    result = subprocess.run(
        ["yt-dlp", "--cookies-from-browser", "chrome",
         "--print", "%(title)s\t%(duration)s", "--no-playlist", url],
        capture_output=True, text=True,
    )
    if result.returncode != 0 or not result.stdout.strip():
        return None
    parts = result.stdout.strip().split("\t")
    if len(parts) != 2:
        return None
    try:
        return {
            "title": sanitize_filename(parts[0].strip()),
            "duration": float(parts[1].strip()),
        }
    except ValueError:
        return None

def download_audio(url: str, output_path: str):
    result = subprocess.run(
        ["yt-dlp", "--cookies-from-browser", "chrome",
         "-x", "--audio-format", "mp3", "--audio-quality", "0",
         "--no-playlist", "-o", output_path, url],
        capture_output=True, text=True,
    )
    if result.returncode != 0:
        return False, result.stderr.strip()
    return True, None


# -- Progress file ------------------------------------------------------------
# The progress file is a plain JSON object: { "completed_urls": [...] }
# It lives next to the spreadsheet and is updated every time a video
# is successfully transcribed (or skipped). On resume, any URL already
# in this set is skipped automatically.

def load_progress(progress_path: str) -> set:
    """Return the set of URLs already completed from a previous run."""
    if not os.path.exists(progress_path):
        return set()
    try:
        with open(progress_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return set(data.get("completed_urls", []))
    except (json.JSONDecodeError, OSError):
        return set()

def save_progress(progress_path: str, completed_urls: set):
    """Atomically write the progress file so a crash never corrupts it."""
    tmp = progress_path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump({"completed_urls": sorted(completed_urls)}, f, indent=2)
    os.replace(tmp, progress_path)  # atomic on Windows and POSIX


# -- Model Loading ------------------------------------------------------------

def load_model(model_size: str, batch_size: int) -> tuple:
    device = "cuda" if torch.cuda.is_available() else "cpu"
    compute_type = "float16" if device == "cuda" else "int8"

    print(f"  Device       : {device.upper()}")
    print(f"  Compute type : {compute_type}")
    print(f"  Batch size   : {batch_size}")

    if device == "cpu":
        print("  WARNING: CUDA not detected -- falling back to CPU. This will be much slower.")
        print("  To fix: pip install torch --index-url https://download.pytorch.org/whl/cu121")

    print(f"  Loading model '{model_size}' (cached after first run)...")
    base_model = WhisperModel(model_size, device=device, compute_type=compute_type)
    pipeline = BatchedInferencePipeline(model=base_model)
    print("  Model ready.\n")
    return pipeline, batch_size


# -- Transcription ------------------------------------------------------------

def transcribe_audio(audio_path: str, pipeline: BatchedInferencePipeline, batch_size: int) -> list[dict]:
    segments, _ = pipeline.transcribe(
        audio_path,
        word_timestamps=True,
        vad_filter=True,
        vad_parameters=dict(
            min_speech_duration_ms=250,
            min_silence_duration_ms=500,
        ),
        batch_size=batch_size,
    )
    sentences_with_times = []
    for segment in segments:
        text = segment.text.strip()
        if not text:
            continue
        sentences = sent_tokenize(text)
        words = list(segment.words) if segment.words else []
        word_idx = 0
        for sentence in sentences:
            word_count = len(sentence.split())
            sentence_words = words[word_idx : word_idx + word_count]
            start = sentence_words[0].start if sentence_words else segment.start
            end = sentence_words[-1].end if sentence_words else segment.end
            sentences_with_times.append({"sentence": sentence, "start": start, "end": end})
            word_idx += word_count
    return sentences_with_times


# -- Output -------------------------------------------------------------------

def save_docx(sentences: list[dict], output_path: str, source_name: str):
    doc = Document()
    doc.add_heading(f"Transcription: {source_name}", level=1)
    for item in sentences:
        start = format_timestamp(item["start"])
        end = format_timestamp(item["end"])
        paragraph = doc.add_paragraph()
        run_ts = paragraph.add_run(f"[{start} -> {end}]  ")
        run_ts.bold = True
        paragraph.add_run(item["sentence"])
    doc.save(output_path)


# -- Logging ------------------------------------------------------------------

_log_lock = threading.Lock()

def write_log(log_path: str, entry: dict):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    entry["logged_at"] = timestamp
    with _log_lock:
        with open(log_path, "a", encoding="utf-8") as f:
            f.write("\n" + "-" * 60 + "\n")
            f.write(f"  Logged at  : {timestamp}\n")
            f.write(f"  Status     : {entry.get('status', 'unknown')}\n")
            f.write(f"  Title      : {entry.get('title', 'N/A')}\n")
            f.write(f"  URL        : {entry.get('url', 'N/A')}\n")
            f.write(f"  Duration   : {entry.get('duration_readable', 'N/A')}\n")
            f.write(f"  Audio file : {entry.get('audio_path', 'N/A')}\n")
            f.write(f"  Transcript : {entry.get('transcript_path', 'N/A')}\n")
            if entry.get("error"):
                f.write(f"  Error      : {entry['error']}\n")
            f.write(f"  JSON       : {json.dumps(entry)}\n")


# -- Pipelined Processing -----------------------------------------------------

_DONE = object()

def _downloader_thread(pending_rows, audio_dir, output_dir, skip_existing, log_path, ready_queue, download_only=False):
    for row_num, url, course_name in pending_rows:
        if _stop.is_set():
            break

        meta = get_video_metadata(url)
        if not meta:
            print(f"  [DL] Row {row_num}: Could not fetch metadata -- {url}")
            write_log(log_path, {
                "url": url, "title": None, "audio_path": None, "transcript_path": None,
                "duration_raw": None, "duration_readable": None,
                "status": "failed", "error": "Could not fetch video metadata",
            })
            ready_queue.put({"job": (row_num, url, None, course_name), "audio_path": None,
                             "log_base": None, "error": "Could not fetch video metadata",
                             "meta_error": True})
            continue

        title             = meta["title"]
        duration          = meta["duration"]
        course_audio_dir  = os.path.join(audio_dir, course_name)
        os.makedirs(course_audio_dir, exist_ok=True)
        audio_path        = os.path.join(course_audio_dir, f"{title}.mp3")

        if skip_existing and not download_only:
            course_output_dir = os.path.join(output_dir, course_name)
            output_path = os.path.join(course_output_dir, f"{title}.docx")
            if os.path.exists(output_path):
                print(f"  [DL] Row {row_num}: {title} -- output already exists, skipping.")
                ready_queue.put({"job": (row_num, url, meta, course_name), "audio_path": None,
                                 "log_base": {
                                     "url": url, "title": title,
                                     "audio_path": str(Path(audio_path).resolve()),
                                     "transcript_path": str(Path(output_path).resolve()),
                                     "duration_raw": duration,
                                     "duration_readable": format_duration(duration),
                                     "status": "skipped", "error": None,
                                 }, "error": None, "skipped": True})
                continue

        log_base = {
            "url": url, "title": title,
            "audio_path": str(Path(audio_path).resolve()),
            "transcript_path": None,
            "duration_raw": duration,
            "duration_readable": format_duration(duration),
            "status": None, "error": None,
        }

        if os.path.exists(audio_path):
            print(f"  [DL] {title} -- audio already on disk.")
            if download_only:
                log_base["status"] = "downloaded"
                write_log(log_path, log_base)
            ready_queue.put({"job": (row_num, url, meta, course_name), "audio_path": audio_path,
                             "log_base": log_base, "error": None})
            continue

        print(f"  [DL] Downloading: {title}  ({format_duration(duration)})")
        success, err = download_audio(url, audio_path)
        if not success:
            print(f"  [DL] ERROR: {title}\n       {err}")
            log_base["status"] = "failed"
            log_base["error"]  = f"Audio download failed: {err}"
            write_log(log_path, log_base)
            ready_queue.put({"job": (row_num, url, meta, course_name), "audio_path": None,
                             "log_base": log_base, "error": err})
        else:
            print(f"  [DL] Ready: {title}")
            if download_only:
                log_base["status"] = "downloaded"
                write_log(log_path, log_base)
            ready_queue.put({"job": (row_num, url, meta, course_name), "audio_path": audio_path,
                             "log_base": log_base, "error": None})

    ready_queue.put(_DONE)


def process_spreadsheet_pipelined(
    xlsx_path, audio_dir, output_dir, pipeline, batch_size,
    skip_existing, log_path, progress_path, prefetch=3, download_only=False
):
    wb = load_workbook(xlsx_path)
    ws = wb.active
    URL_COL    = 5  # Column F (zero-based)
    COURSE_COL = 4  # Column E (zero-based)

    # Load previously completed URLs
    completed_urls = load_progress(progress_path)
    if completed_urls:
        print(f"  Resuming -- {len(completed_urls)} URL(s) already completed (from progress file).\n")

    print("Scanning spreadsheet...")
    failed = []

    raw_rows = [
        (
            i + 1,
            str(row[URL_COL].value).strip(),
            sanitize_filename(str(row[COURSE_COL].value).strip()) if row[COURSE_COL].value else "Uncategorised",
        )
        for i, row in enumerate(ws.iter_rows())
        if row[URL_COL].value and str(row[URL_COL].value).strip().startswith("http")
    ]

    if not raw_rows:
        print("No URLs found in column F of the spreadsheet.")
        return

    total_rows = len(raw_rows)
    print(f"Found {total_rows} URL(s).\n")

    pending_rows = []
    for row_num, url, course_name in raw_rows:
        if url in completed_urls:
            print(f"  [RESUME] Row {row_num}: already completed -- skipping.")
            continue
        pending_rows.append((row_num, url, course_name))

    if not pending_rows:
        print("\nNothing left to process (all URLs already completed).")
        _print_summary(total_rows, failed, log_path, progress_path)
        return

    print(f"\n{len(pending_rows)} URL(s) to process.")
    if download_only:
        print("Starting download-only run (no transcription)...")
    else:
        print("Starting pipelined download + transcription...")
    print("  [DL] = downloader thread      [TX] = transcriber (GPU)")
    print("  Press CTRL+C once to stop gracefully after current video.\n")

    ready_queue = queue.Queue(maxsize=prefetch)
    dl_thread = threading.Thread(
        target=_downloader_thread,
        args=(pending_rows, audio_dir, output_dir, skip_existing, log_path, ready_queue, download_only),
        daemon=True,
    )
    dl_thread.start()

    completed = 0
    stopped_early = False

    while True:
        item = ready_queue.get()
        if item is _DONE:
            break

        if _stop.is_set():
            print("\n  [!] Stopping after current item. Run the same command to resume.")
            stopped_early = True
            break

        row_num, url, meta, course_name = item["job"]

        if item.get("meta_error"):
            failed.append((row_num, url, "Could not fetch video metadata"))
            continue

        if item.get("skipped"):
            write_log(log_path, item["log_base"])
            completed_urls.add(url)
            save_progress(progress_path, completed_urls)
            continue

        audio_path = item["audio_path"]
        log_base   = item["log_base"]
        dl_error   = item["error"]
        title      = meta["title"]

        course_output_dir = os.path.join(output_dir, course_name)
        os.makedirs(course_output_dir, exist_ok=True)
        output_path = os.path.join(course_output_dir, f"{title}.docx")

        log_base["transcript_path"] = str(Path(output_path).resolve())

        completed += 1

        if download_only:
            if dl_error:
                failed.append((row_num, url, "Download failed"))
                continue
            # Log already written by downloader; just mark progress
            completed_urls.add(url)
            save_progress(progress_path, completed_urls)
            print(f"  [DL] ({completed}) Downloaded: {title}\n")
            continue

        print(f"  [TX] ({completed}) Transcribing: {title}")

        if dl_error:
            failed.append((row_num, url, "Download failed"))
            continue

        sentences = transcribe_audio(audio_path, pipeline, batch_size)
        if not sentences:
            print(f"  [TX] WARNING: Empty transcription for {title}")
            log_base["status"] = "failed"
            log_base["error"]  = "Empty transcription result"
            write_log(log_path, log_base)
            failed.append((row_num, url, "Empty transcription"))
            continue

        save_docx(sentences, output_path, title)
        log_base["status"] = "success"
        write_log(log_path, log_base)

        # Mark this URL as complete in the progress file immediately
        completed_urls.add(url)
        save_progress(progress_path, completed_urls)

        print(f"  [TX] Done: {output_path}\n")

    dl_thread.join(timeout=2)

    if stopped_early:
        print(f"  Stopped early.")
        print(f"  Progress saved to: {progress_path}")
        print(f"  Run the same command again to resume from where you left off.\n")

    _print_summary(total_rows, failed, log_path, progress_path)


def _print_summary(total, failed, log_path, progress_path=None):
    print("\n" + "=" * 50)
    print(f"Completed: {total - len(failed)}/{total}")
    if failed:
        print(f"\nFailed ({len(failed)}):")
        for row_num, url, reason in failed:
            print(f"  Row {row_num}: {reason} -- {url}")
    else:
        print("All rows completed successfully.")
    print("=" * 50)
    print(f"\nLog written to   : {log_path}")
    if progress_path:
        print(f"Progress file    : {progress_path}")


# -- Directory mode -----------------------------------------------------------

def process_directory(audio_dir, output_dir, pipeline, batch_size, skip_existing, log_path):
    extensions = {".mp3", ".mp4", ".wav", ".m4a", ".ogg", ".flac", ".webm"}
    files = sorted([f for f in Path(audio_dir).iterdir()
                    if f.is_file() and f.suffix.lower() in extensions])
    if not files:
        print(f"No audio files found in: {audio_dir}")
        return
    total = len(files)
    print(f"\nFound {total} audio file(s). Starting...\n")
    failed = []
    for i, audio_file in enumerate(files, 1):
        if _stop.is_set():
            print("\n  [!] Stopping. Re-run the same command to continue from where you left off.")
            break
        stem        = audio_file.stem
        output_path = os.path.join(output_dir, f"{stem}.docx")
        print(f"[{i}/{total}] {audio_file.name}")
        log_entry = {
            "url": "N/A", "title": stem,
            "audio_path": str(audio_file.resolve()),
            "transcript_path": str(Path(output_path).resolve()),
            "duration_raw": None, "duration_readable": "N/A",
            "status": None, "error": None,
        }
        if skip_existing and os.path.exists(output_path):
            print(f"    Skipping -- transcript already exists.")
            log_entry["status"] = "skipped"
            write_log(log_path, log_entry)
            print()
            continue
        sentences = transcribe_audio(str(audio_file), pipeline, batch_size)
        if not sentences:
            print(f"    WARNING: Empty transcription.")
            log_entry["status"] = "failed"
            log_entry["error"]  = "Empty transcription result"
            write_log(log_path, log_entry)
            failed.append(audio_file.name)
            print()
            continue
        save_docx(sentences, output_path, stem)
        log_entry["status"] = "success"
        write_log(log_path, log_entry)
        print(f"    Saved: {output_path}\n")
    print("=" * 50)
    print(f"Completed: {total - len(failed)}/{total} files")
    if failed:
        print(f"\nFailed files:")
        for name in failed:
            print(f"  {name}")
    else:
        print("All files completed successfully.")
    print("=" * 50)
    print(f"\nLog written to: {log_path}")


# -- CLI Entry Point ----------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="YouTube Audio Transcriber -> DOCX with timestamps (faster-whisper + CUDA, pipelined, resumable)"
    )
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--spreadsheet", "-s", help="Path to .xlsx file with YouTube URLs in column F.")
    group.add_argument("--directory",   "-d", help="Path to a directory of audio files to transcribe.")
    parser.add_argument("--audio-dir",  "-a", default="./audio",
                        help="Where to save downloaded audio. (default: ./audio)")
    parser.add_argument("--output-dir", "-o", default="./transcripts",
                        help="Where to save .docx files. (default: ./transcripts)")
    parser.add_argument("--log-file",   "-l", default="./transcription_log.txt",
                        help="Path to the log file. (default: ./transcription_log.txt)")
    parser.add_argument("--progress-file", default=None,
                        help="Path to the progress/resume JSON file. "
                             "Defaults to <spreadsheet>.progress.json next to the spreadsheet.")
    parser.add_argument("--model", "-m", default="large-v2",
                        choices=["tiny", "base", "small", "medium", "large-v1", "large-v2", "large-v3"],
                        help="Whisper model size. (default: large-v2)")
    parser.add_argument("--batch-size", "-b", type=int, default=16,
                        help="GPU chunks processed in parallel. Reduce if out of VRAM. (default: 16)")
    parser.add_argument("--prefetch", "-p", type=int, default=3,
                        help="Audio files to download ahead of transcription. (default: 3)")
    parser.add_argument("--no-skip", action="store_true",
                        help="Re-transcribe even if a .docx already exists.")
    parser.add_argument("--reset-progress", action="store_true",
                        help="Delete the progress file and start from scratch.")
    parser.add_argument("--download-only", action="store_true",
                        help="Download audio only -- do not load Whisper or transcribe. "
                             "Only valid with --spreadsheet.")

    args = parser.parse_args()
    os.makedirs(args.audio_dir,  exist_ok=True)
    os.makedirs(args.output_dir, exist_ok=True)

    # Determine progress file path
    if args.progress_file:
        progress_path = args.progress_file
    elif args.spreadsheet:
        progress_path = str(Path(args.spreadsheet).with_suffix(".progress.json"))
    else:
        progress_path = "./transcription.progress.json"

    if args.reset_progress and os.path.exists(progress_path):
        os.remove(progress_path)
        print(f"  Progress file deleted -- starting from scratch.\n")

    with open(args.log_file, "a", encoding="utf-8") as f:
        f.write("\n" + "=" * 60 + "\n")
        f.write(f"  SESSION START: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        if args.spreadsheet:
            mode_label = "download-only" if args.download_only else "spreadsheet"
        else:
            mode_label = "directory"
        f.write(f"  Mode         : {mode_label}\n")
        f.write(f"  Input        : {args.spreadsheet or args.directory}\n")
        f.write(f"  Audio dir    : {Path(args.audio_dir).resolve()}\n")
        f.write(f"  Output dir   : {Path(args.output_dir).resolve()}\n")
        if not args.download_only:
            f.write(f"  Whisper model: {args.model}\n")
            f.write(f"  Batch size   : {args.batch_size}\n")
        else:
            f.write(f"  Whisper model: N/A\n")
            f.write(f"  Batch size   : N/A\n")
        f.write(f"  Prefetch     : {args.prefetch}\n")
        f.write(f"  Skip existing: {not args.no_skip}\n")
        f.write(f"  Progress file: {progress_path}\n")
        f.write("=" * 60 + "\n")

    if args.download_only and not args.spreadsheet:
        parser.error("--download-only requires --spreadsheet")

    if not args.download_only:
        print(f"Loading faster-whisper model '{args.model}'...")
        pipeline, batch_size = load_model(args.model, args.batch_size)
    else:
        pipeline, batch_size = None, None

    if args.spreadsheet:
        process_spreadsheet_pipelined(
            args.spreadsheet, args.audio_dir, args.output_dir,
            pipeline, batch_size, not args.no_skip, args.log_file,
            progress_path, prefetch=args.prefetch, download_only=args.download_only,
        )
    elif args.directory:
        process_directory(
            args.directory, args.output_dir,
            pipeline, batch_size, not args.no_skip, args.log_file,
        )


if __name__ == "__main__":
    main()
